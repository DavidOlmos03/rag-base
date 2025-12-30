"""
DOCX document parser implementation.
"""

import os

from docx import Document

from app.core.exceptions import DocumentParsingError
from app.core.logging import get_logger
from app.domain.interfaces.document_parser import DocumentParser, ParsedDocument

logger = get_logger(__name__)


class DOCXParser(DocumentParser):
    """DOCX parser using python-docx."""

    async def parse(self, file_path: str) -> ParsedDocument:
        """Parse a DOCX document and extract text.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Parsed document with content and metadata

        Raises:
            DocumentParsingError: If parsing fails
        """
        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_content = "\n\n".join(paragraphs)

            # Extract metadata from core properties
            metadata = {}
            if doc.core_properties:
                metadata = {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "keywords": doc.core_properties.keywords or "",
                    "comments": doc.core_properties.comments or "",
                    "category": doc.core_properties.category or "",
                }

            # Get file stats
            file_size = os.path.getsize(file_path)

            # Count sections/pages (approximate)
            num_sections = len(doc.sections)

            logger.info(
                "docx_parsed",
                file_path=file_path,
                num_paragraphs=len(paragraphs),
                content_length=len(full_content),
            )

            return ParsedDocument(
                content=full_content,
                metadata=metadata,
                num_pages=num_sections,
                file_type="docx",
                file_size=file_size,
            )

        except Exception as e:
            logger.error("docx_parsing_failed", file_path=file_path, error=str(e))
            raise DocumentParsingError(f"Failed to parse DOCX: {str(e)}") from e

    def supports_format(self, file_extension: str) -> bool:
        """Check if parser supports file format.

        Args:
            file_extension: File extension (e.g., '.docx')

        Returns:
            True if format is supported
        """
        return file_extension.lower() in self.supported_formats

    @property
    def supported_formats(self) -> list[str]:
        """Get list of supported file formats.

        Returns:
            List of supported extensions
        """
        return [".docx", ".doc"]
